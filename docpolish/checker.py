import os
from docx import Document
from docx.shared import RGBColor
import language_tool_python
import difflib

# === Initialize Tools ===
print("⏳ Loading grammar tool...")
tool = language_tool_python.LanguageTool('en-US')
print("✅ Tool loaded successfully.")

# === Highlight differences between original and corrected ===
def highlight_differences(original, corrected):
    diff = difflib.ndiff(original.split(), corrected.split())
    highlighted = []
    for d in diff:
        code = d[0]
        word = d[2:]
        if code == "-":
            highlighted.append(f"**{word}**")
        elif code == " ":
            highlighted.append(word)
    return " ".join(highlighted)

# === Correct entire paragraph ===
def correct_paragraph(paragraph):
    matches = tool.check(paragraph)
    if not matches:
        return {
            'original': paragraph,
            'corrected': paragraph,
            'highlighted': paragraph,
            'needs_correction': False
        }

    corrected = language_tool_python.utils.correct(paragraph, matches)
    highlighted = highlight_differences(paragraph, corrected)
    return {
        'original': paragraph,
        'corrected': corrected,
        'highlighted': highlighted,
        'needs_correction': True
    }

# === Process DOCX paragraph-wise ===
def process_docx_paragraphs(input_path, output_path):
    print(f"🔍 Processing document paragraph-wise: {input_path}")
    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        result = correct_paragraph(para.text)
        new_para = new_doc.add_paragraph()

        if result['needs_correction']:
            run = new_para.add_run(result['highlighted'] + " ")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            run.italic = True

            suggestion = new_para.add_run(f"(Suggestion: {result['corrected']}) ")
            suggestion.italic = True
            suggestion.font.color.rgb = RGBColor(0, 128, 0)
        else:
            new_para.add_run(result['original'])

    new_doc.save(output_path)
    print(f"✅ Polished document saved: {output_path}")

# === Run ===
if __name__ == "__main__":
    input_path = input("📥 Enter the full path to the input .docx file: ").strip()
    output_path = input("📤 Enter the desired path to save the polished .docx file: ").strip()
    process_docx_paragraphs(input_path, output_path)
